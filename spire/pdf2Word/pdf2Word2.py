from spire.pdf.common import *
from spire.pdf import *
import os
from docx import Document

input_pdf = "Sample.pdf"
output_word = "output.docx"

def extract_text_from_pdf(pdf_path):
    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")

    doc = PdfDocument()
    doc.LoadFromFile(pdf_path)

    all_pages_text = []
    for i in range(doc.Pages.Count):
        page = doc.Pages[i]
        extractor = PdfTextExtractor(page)
        options = PdfTextExtractOptions()
        text = extractor.ExtractText(options)
        all_pages_text.append(text)

    doc.Close()
    return all_pages_text

def create_word_doc(text_pages, filename):
    doc = Document()
    for i, page_text in enumerate(text_pages, 1):
        doc.add_heading(f"Page {i}", level=2)
        doc.add_paragraph(page_text)
        doc.add_page_break()
    doc.save(filename)
    print(f"Word file saved: {filename}")

def main():
    try:
        pages_text = extract_text_from_pdf(input_pdf)
        create_word_doc(pages_text, output_word)
        print("Word document created successfully!")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()
